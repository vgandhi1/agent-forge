"""Generate a sample NDA .docx to give the POC something real to revise."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def make_sample_nda(path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=0)

    doc.add_paragraph(
        'This Mutual Non-Disclosure Agreement (the "Agreement") is entered into as of '
        'January 1, 2026 (the "Effective Date") by and between Acme Corp. and Beta LLC '
        '(each a "Party" and collectively the "Parties").'
    )
    doc.add_heading("1. Definition of Confidential Information", level=1)
    doc.add_paragraph(
        '"Confidential Information" means any non-public information disclosed by one Party '
        "to the other, whether orally or in writing, that is designated as confidential."
    )
    doc.add_heading("2. Term", level=1)
    doc.add_paragraph(
        "This Agreement shall remain in effect for a period of two (2) years from the "
        "Effective Date, unless terminated earlier by either Party with thirty (30) days "
        "written notice."
    )
    doc.add_heading("3. Obligations", level=1)
    doc.add_paragraph(
        "Each Party agrees to hold the other Party's Confidential Information in strict "
        "confidence and not to disclose it to any third party without prior written consent."
    )
    doc.add_heading("4. Governing Law", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with the laws of "
        "the State of California."
    )
    doc.save(str(path))
    return path


if __name__ == "__main__":
    import sys

    out = sys.argv[1] if len(sys.argv) > 1 else "contracts/workspace/sample_nda.docx"
    p = make_sample_nda(out)
    print(f"Wrote sample NDA → {p}")
